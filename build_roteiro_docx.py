import json
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Desafio_Cafe_Brasil_Roteiro_Questoes.docx"
SCRIPT = "script.js"


def load_questions():
    js = r"""
const fs = require("fs");
const source = fs.readFileSync("script.js", "utf8");
const start = source.indexOf("const questions = ");
const end = source.indexOf(";\n\nlet order", start);
if (start < 0 || end < 0) {
  throw new Error("Não foi possível localizar o array questions em script.js");
}
const body = source.slice(start + "const questions = ".length, end);
const questions = Function(`return (${body});`)();
console.log(JSON.stringify(questions));
"""
    result = subprocess.run(
        ["node", "-e", js],
        check=True,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )
    return json.loads(result.stdout)


def set_styles(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.color.rgb = RGBColor(31, 77, 120)

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_label(document, label, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(text or "")


def add_link_label(document, label, item):
    if not item:
        add_label(
            document,
            label,
            "não selecionada nesta questão; usar somente publicação/ATER+Digital ou realizar curadoria posterior.",
        )
        return
    add_label(document, label, item["title"])
    add_label(document, "Link", item["url"])


def add_question(document, index, question):
    document.add_heading(f"Questão {index}", level=3)
    add_label(document, "Bloco", question["block"])
    add_label(document, "Categoria", question["category"])
    add_label(document, "Enunciado", question["question"])

    for i, option in enumerate(question["options"]):
        letter = chr(65 + i)
        document.add_paragraph(f"{letter}) {option}", style="List Bullet")

    correct_letter = chr(65 + question["answer"])
    add_label(document, "Resposta correta", f"{correct_letter}) {question['options'][question['answer']]}")
    add_label(document, "Explicação curta", question["explanation"])

    document.add_paragraph("Conexão Embrapa").runs[0].bold = True
    connection = question["connection"]
    add_link_label(document, "ATER+Digital", connection.get("ater"))
    add_link_label(document, "Publicação relacionada", connection.get("publication"))
    add_link_label(document, "Solução tecnológica relacionada", connection.get("solution"))


def main():
    if not Path(SCRIPT).exists():
        raise FileNotFoundError(SCRIPT)

    questions = load_questions()
    document = Document()
    set_styles(document)

    document.add_paragraph("Quiz Interativo - Sistemas Cafeeiros e Conexão Embrapa", style="Title")
    document.add_paragraph(
        "Versão estruturada com questões, alternativas completas, resposta correta e conexão com ATER+Digital, publicações e soluções tecnológicas da Embrapa."
    )

    document.add_heading("Objetivo editorial", level=1)
    document.add_paragraph(
        "Substituir a lógica de disputa entre Arábica e Canéfora por uma abordagem integrada de sistemas cafeeiros, destacando diversidade, território, genética, manejo, pós-colheita, qualidade, sustentabilidade e inovação."
    )

    document.add_heading("Fontes usadas para curadoria", level=1)
    for item in [
        "Conteúdo ATER+Digital Café enviado pela equipe.",
        "Planilha publicacoes_cafe_clusteres.xlsx.",
        "Planilha solucoes_cafe2.xlsx.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    current_block = None
    for index, question in enumerate(questions, start=1):
        if question["block"] != current_block:
            current_block = question["block"]
            document.add_heading(current_block, level=2)
        add_question(document, index, question)

    document.add_heading("Observação editorial", level=1)
    document.add_paragraph(
        "As conexões com a Embrapa foram redigidas como indicações de aprofundamento. Na implementação web, elas podem ser ligadas à curadoria final do acervo (cafe > autores > BigQuery)."
    )

    document.save(OUTPUT)


if __name__ == "__main__":
    main()
